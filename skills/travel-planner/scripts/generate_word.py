"""
Travel Planner - Word Document Generator
Generates professional Word travel guide documents.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path


def create_word_document(travel_data: dict, output_path: str = None) -> str:
    """
    Generate a Word document from travel planning data.

    Args:
        travel_data: Dictionary containing all travel information
        output_path: Optional output file path

    Returns:
        Path to generated file
    """
    doc = Document()

    # Title
    title = doc.add_heading(travel_data.get("title", "旅游攻略"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with basic info
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{travel_data.get('destination', '')} | "
                    f"{travel_data.get('duration', '')}天\n")
    subtitle.add_run(f"出行时间: {travel_data.get('start_date', '')} - "
                    f"{travel_data.get('end_date', '')}")

    # Cover image if available
    if travel_data.get("cover_image"):
        doc.add_picture(travel_data["cover_image"], width=Inches(6))

    # Table of Contents
    doc.add_heading("目录", level=1)
    toc = doc.add_paragraph()
    toc.add_run("1. 行程概览\n")
    toc.add_run("2. 目的地信息\n")
    toc.add_run("3. 详细行程\n")
    toc.add_run("4. 住宿推荐\n")
    toc.add_run("5. 美食推荐\n")
    toc.add_run("6. 交通指南\n")
    toc.add_run("7. 预算规划\n")
    toc.add_run("8. 注意事项\n")
    toc.add_run("9. 天气预报\n")

    # Section 1: Overview
    doc.add_heading("1. 行程概览", level=1)
    overview = doc.add_paragraph()
    overview.add_run(f"出行天数: ").bold = True
    overview.add_run(f"{travel_data.get('duration', 'N/A')}天\n")
    overview.add_run(f"同行人员: ").bold = True
    overview.add_run(f"{travel_data.get('companions', 'N/A')}\n")
    overview.add_run(f"总预算: ").bold = True
    overview.add_run(f"{travel_data.get('budget', 'N/A')}元\n")
    overview.add_run(f"旅行主题: ").bold = True
    overview.add_run(f"{travel_data.get('theme', 'N/A')}")

    # Section 2: Destination Info
    doc.add_heading("2. 目的地信息", level=1)
    dest_info = travel_data.get("destination_info", {})
    if dest_info:
        p = doc.add_paragraph()
        p.add_run(dest_info.get("description", ""))
    else:
        doc.add_paragraph("目的地信息加载中...")

    # Section 3: Detailed Itinerary
    doc.add_heading("3. 详细行程", level=1)
    itinerary = travel_data.get("itinerary", [])
    if itinerary:
        for i, day in enumerate(itinerary, 1):
            doc.add_heading(f"第{i}天: {day.get('date', f'Day {i}')}", level=2)
            day_plan = doc.add_paragraph()

            # Morning
            morning = day.get("morning", {})
            if morning:
                day_plan.add_run("上午: ").bold = True
                day_plan.add_run(f"{morning.get('activity', '')} "
                                f"({morning.get('time', '')}, "
                                f"地址: {morning.get('address', 'N/A')})\n")

            # Afternoon
            afternoon = day.get("afternoon", {})
            if afternoon:
                day_plan.add_run("下午: ").bold = True
                day_plan.add_run(f"{afternoon.get('activity', '')} "
                                f"({afternoon.get('time', '')}, "
                                f"地址: {afternoon.get('address', 'N/A')})\n")

            # Evening
            evening = day.get("evening", {})
            if evening:
                day_plan.add_run("晚上: ").bold = True
                day_plan.add_run(f"{evening.get('activity', '')} "
                                f"({evening.get('time', '')}, "
                                f"地址: {evening.get('address', 'N/A')})\n")

            # Meals
            meals = day.get("meals", {})
            if meals:
                day_plan.add_run("用餐: ").bold = True
                day_plan.add_run(f"早餐: {meals.get('breakfast', '自理')} | "
                                f"午餐: {meals.get('lunch', '自理')} | "
                                f"晚餐: {meals.get('dinner', '自理')}")

            # Transportation
            transport = day.get("transport", {})
            if transport:
                doc.add_paragraph()
                transport_p = doc.add_paragraph()
                transport_p.add_run("交通: ").bold = True
                transport_p.add_run(f"{transport.get('method', '')} "
                                    f"({transport.get('details', '')})")
    else:
        doc.add_paragraph("行程规划中...")

    # Section 4: Hotels
    doc.add_heading("4. 住宿推荐", level=1)
    hotels = travel_data.get("hotels", [])
    if hotels:
        for i, hotel in enumerate(hotels, 1):
            doc.add_heading(f"推荐酒店 {i}", level=3)
            hotel_p = doc.add_paragraph()
            hotel_p.add_run(f"名称: ").bold = True
            hotel_p.add_run(f"{hotel.get('name', 'N/A')}\n")
            hotel_p.add_run(f"地址: ").bold = True
            hotel_p.add_run(f"{hotel.get('address', 'N/A')}\n")
            hotel_p.add_run(f"价格: ").bold = True
            hotel_p.add_run(f"{hotel.get('price', 'N/A')}/晚\n")
            hotel_p.add_run(f"评分: ").bold = True
            hotel_p.add_run(f"{hotel.get('rating', 'N/A')}分\n")
            hotel_p.add_run(f"特色: ").bold = True
            hotel_p.add_run(f"{hotel.get('features', 'N/A')}")
    else:
        doc.add_paragraph("住宿规划中...")

    # Section 5: Food
    doc.add_heading("5. 美食推荐", level=1)
    foods = travel_data.get("food_recommendations", [])
    if foods:
        for i, food in enumerate(foods, 1):
            doc.add_heading(f"美食 {i}: {food.get('name', 'N/A')}", level=3)
            food_p = doc.add_paragraph()
            food_p.add_run(f"类型: ").bold = True
            food_p.add_run(f"{food.get('type', 'N/A')}\n")
            food_p.add_run(f"地址: ").bold = True
            food_p.add_run(f"{food.get('address', 'N/A')}\n")
            food_p.add_run(f"人均: ").bold = True
            food_p.add_run(f"{food.get('price_per_person', 'N/A')}元\n")
            food_p.add_run(f"推荐菜: ").bold = True
            food_p.add_run(f"{food.get('recommended_dishes', 'N/A')}")
    else:
        doc.add_paragraph("美食推荐规划中...")

    # Section 6: Transportation
    doc.add_heading("6. 交通指南", level=1)
    transport_info = travel_data.get("transport_info", {})
    if transport_info:
        # Getting there
        getting_there = transport_info.get("getting_there", {})
        doc.add_heading("到达方式", level=2)
        gt_p = doc.add_paragraph()
        gt_p.add_run(f"方式: ").bold = True
        gt_p.add_run(f"{getting_there.get('method', 'N/A')}\n")
        gt_p.add_run(f"详情: ").bold = True
        gt_p.add_run(f"{getting_there.get('details', 'N/A')}\n")
        gt_p.add_run(f"预计时间: ").bold = True
        gt_p.add_run(f"{getting_there.get('duration', 'N/A')}\n")
        gt_p.add_run(f"费用: ").bold = True
        gt_p.add_run(f"{getting_there.get('cost', 'N/A')}元")

        # Local transport
        doc.add_heading("当地交通", level=2)
        local = transport_info.get("local_transport", [])
        if local:
            for item in local:
                lp = doc.add_paragraph()
                lp.add_run(f"- {item.get('type', '')}: {item.get('description', '')}")
        else:
            doc.add_paragraph("当地交通规划中...")
    else:
        doc.add_paragraph("交通指南规划中...")

    # Section 7: Budget
    doc.add_heading("7. 预算规划", level=1)
    budget = travel_data.get("budget_breakdown", {})
    if budget:
        budget_p = doc.add_paragraph()
        budget_p.add_run(f"总预算: ").bold = True
        budget_p.add_run(f"{budget.get('total', 'N/A')}元\n\n")

        items = [
            ("交通", budget.get("transport", 0)),
            ("住宿", budget.get("accommodation", 0)),
            ("餐饮", budget.get("food", 0)),
            ("门票", budget.get("tickets", 0)),
            ("购物", budget.get("shopping", 0)),
            ("其他", budget.get("other", 0)),
        ]

        for category, amount in items:
            bp = doc.add_paragraph()
            bp.add_run(f"{category}: ").bold = True
            bp.add_run(f"{amount}元")

        # Actual vs budget comparison table
        doc.add_paragraph()
        doc.add_heading("预算对比表", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"

        # Headers
        headers = table.rows[0].cells
        headers[0].text = "类别"
        headers[1].text = "预算"
        headers[2].text = "实际花费"

        # Sample data rows
        for row_idx, (category, budget_amount) in enumerate([
            ("交通", budget.get("transport", 0)),
            ("住宿", budget.get("accommodation", 0)),
            ("餐饮", budget.get("food", 0)),
        ], 1):
            row = table.rows[row_idx].cells
            row[0].text = category
            row[1].text = f"{budget_amount}元"
            row[2].text = "待填写"
    else:
        doc.add_paragraph("预算规划中...")

    # Section 8: Notes
    doc.add_heading("8. 注意事项", level=1)
    notes = travel_data.get("notes", [])
    if notes:
        for note in notes:
            doc.add_paragraph(f"- {note}", style="List Bullet")
    else:
        doc.add_paragraph("注意事项规划中...")

    # Section 9: Weather
    doc.add_heading("9. 天气预报", level=1)
    weather = travel_data.get("weather", {})
    if weather:
        weather_p = doc.add_paragraph()
        weather_p.add_run(f"目的地: ").bold = True
        weather_p.add_run(f"{weather.get('city', 'N/A')}\n")
        if weather.get("current"):
            current = weather["current"]
            weather_p.add_run(f"当前天气: ").bold = True
            weather_p.add_run(f"{current.get('weather', 'N/A')}, "
                            f"{current.get('temperature', 'N/A')}°C\n")

        forecast = weather.get("forecast", [])
        if forecast:
            doc.add_paragraph()
            weather_p.add_run("天气预报: ").bold = True
            for day in forecast:
                weather_p.add_run(f"\n{day.get('date', '')}: "
                                f"{day.get('day_weather', '')}, "
                                f"{day.get('night_temp', 'N/A')}°C ~ "
                                f"{day.get('day_temp', 'N/A')}°C")
    else:
        doc.add_paragraph("天气信息获取中...")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
                  f"Travel Planner").italic = True

    # Save
    if not output_path:
        destination = travel_data.get("destination", "travel").replace(" ", "_")
        start_date = travel_data.get("start_date", "").replace("-", "")
        output_path = f"travel_guide_{destination}_{start_date}.docx"

    doc.save(output_path)
    return output_path


def generate_sample_travel_data() -> dict:
    """Generate sample travel data for testing."""
    return {
        "title": "杭州旅游攻略",
        "destination": "杭州",
        "duration": 3,
        "start_date": "2024-03-15",
        "end_date": "2024-03-17",
        "companions": "2人（情侣）",
        "budget": 3000,
        "theme": "休闲文化游",
        "destination_info": {
            "description": "杭州是中国著名的历史文化名城，以西湖为核心的风景名胜区闻名遐迩。"
        },
        "itinerary": [
            {
                "date": "2024-03-15",
                "morning": {"activity": "西湖游船", "time": "9:00-12:00", "address": "西湖游船码头"},
                "afternoon": {"activity": "灵隐寺", "time": "14:00-17:00", "address": "灵隐路法云弄1号"},
                "evening": {"activity": "河坊街", "time": "18:30-21:00", "address": "河坊街"},
                "meals": {"breakfast": "酒店早餐", "lunch": "灵隐寺素斋", "dinner": "知味观"}
            },
            {
                "date": "2024-03-16",
                "morning": {"activity": "雷峰塔", "time": "9:00-11:00", "address": "南山路15号"},
                "afternoon": {"activity": "龙井问茶", "time": "14:00-16:00", "address": "龙井村"},
                "evening": {"activity": "宋城千古情", "time": "19:00-21:00", "address": "宋城景区"}
            },
            {
                "date": "2024-03-17",
                "morning": {"activity": "西溪湿地", "time": "9:00-12:00", "address": "文二路西溪湿地北门"},
                "afternoon": {"activity": "购物 & 返程", "time": "14:00-17:00", "address": "银泰百货"}
            }
        ],
        "hotels": [
            {
                "name": "杭州西子湖四季酒店",
                "address": "西湖区灵隐路5号",
                "price": 1800,
                "rating": 4.9,
                "features": "西湖景区内、独栋别墅、管家服务"
            },
            {
                "name": "如家精选酒店(西湖断桥店)",
                "address": "西湖区环城西路12号",
                "price": 350,
                "rating": 4.5,
                "features": "位置好、性价比高、近地铁"
            }
        ],
        "food_recommendations": [
            {
                "name": "外婆家",
                "type": "杭帮菜",
                "address": "西湖区外婆家各分店",
                "price_per_person": 80,
                "recommended_dishes": "西湖醋鱼、东坡肉、龙井虾仁"
            },
            {
                "name": "知味观",
                "type": "老字号",
                "address": "上城区知味观总店",
                "price_per_person": 100,
                "recommended_dishes": "小笼包、猫耳朵、片儿川"
            },
            {
                "name": "新白鹿餐厅",
                "type": "杭帮菜",
                "address": "下城区新白鹿各分店",
                "price_per_person": 60,
                "recommended_dishes": "蛋黄鸡翅、糖醋排骨"
            }
        ],
        "transport_info": {
            "getting_there": {
                "method": "高铁",
                "details": "上海虹桥站 → 杭州东站",
                "duration": "45分钟",
                "cost": 73
            },
            "local_transport": [
                {"type": "地铁", "description": "支持支付宝二维码"},
                {"type": "公交", "description": "2元起，可扫码支付"},
                {"type": "出租车", "description": "11元起步"},
                {"type": "共享单车", "description": "哈啰单车、美团单车"}
            ]
        },
        "budget_breakdown": {
            "total": 3000,
            "transport": 500,
            "accommodation": 1200,
            "food": 600,
            "tickets": 400,
            "shopping": 200,
            "other": 100
        },
        "notes": [
            "西湖游船建议提前在网上购票",
            "灵隐寺门票包含飞来峰，建议早上去",
            "龙井村采茶体验需提前预约",
            "3月是杭州旅游旺季，酒店需提前预订"
        ],
        "weather": {
            "city": "杭州",
            "current": {
                "weather": "多云",
                "temperature": 15,
                "humidity": 60,
                "wind_direction": "东风",
                "wind_power": "3"
            },
            "forecast": [
                {"date": "2024-03-15", "week": "周五", "day_weather": "多云", "night_weather": "晴", "day_temp": "18", "night_temp": "10"},
                {"date": "2024-03-16", "week": "周六", "day_weather": "阴", "night_weather": "小雨", "day_temp": "16", "night_temp": "12"},
                {"date": "2024-03-17", "week": "周日", "day_weather": "小雨", "night_weather": "阴", "day_temp": "14", "night_temp": "11"}
            ]
        }
    }


if __name__ == "__main__":
    sample_data = generate_sample_travel_data()
    output_file = create_word_document(sample_data)
    print(f"Generated: {output_file}")
